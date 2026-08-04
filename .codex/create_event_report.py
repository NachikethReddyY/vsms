from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/nr/developer/react-nodejs-project2-cryptix")
IMAGES = ROOT / ".codex/images"
OUTPUT = ROOT / ".codex/VSMS-Event-Workspace-Implementation-Report.docx"
INK = "181817"
MUTED = "696861"
PAPER = "F1F0EB"
ORANGE = "DF7140"
BLUE = "2879D5"
WHITE = "F7F6F0"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(table, color="D6D4CC", size="6"):
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.find(qn("w:tblBorders"))
    if node is None:
        node = OxmlElement("w:tblBorders")
        tbl_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)
        node.append(tag)


def keep(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("VSMS EVENT WORKSPACE   •   ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def heading(doc, title, kicker=None):
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(kicker.upper())
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(ORANGE)
    h = doc.add_heading(title, level=1)
    h.paragraph_format.space_after = Pt(9)
    return h


def body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def picture(doc, path, caption, width=6.55):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    for run in c.runs:
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    for idx, header in enumerate(headers):
        cell = t.rows[0].cells[idx]
        shade(cell, INK)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(header)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        if widths:
            cell.width = Inches(widths[idx])
    keep(t.rows[0])
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    t.rows[0]._tr.get_or_add_trPr().append(table_header)
    for row_idx, values in enumerate(rows):
        row = t.add_row()
        keep(row)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_idx % 2:
                shade(cell, "F8F7F3")
            cell.paragraphs[0].add_run(str(value))
            if widths:
                cell.width = Inches(widths[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def mobile_pair(doc, left, left_caption, right, right_caption):
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for idx, path in enumerate((left, right)):
        p = t.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(2.28))
    for idx, caption in enumerate((left_caption, right_caption)):
        p = t.cell(1, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    for row in t.rows:
        keep(row)
        for cell in row.cells:
            cell.width = Inches(3.25)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.62)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.header_distance = Inches(0.25)
section.footer_distance = Inches(0.25)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
styles["Normal"].paragraph_format.line_spacing = 1.08
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(34)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor.from_string(INK)
styles["Heading 1"].font.name = "Aptos Display"
styles["Heading 1"].font.size = Pt(23)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor.from_string(INK)
styles["Heading 1"].paragraph_format.keep_with_next = False
styles["Heading 2"].font.name = "Aptos"
styles["Heading 2"].font.size = Pt(14)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor.from_string(INK)

for sec in doc.sections:
    header = sec.header.paragraphs[0]
    header.text = "VSMS  /  DESIGN & IMPLEMENTATION RECORD"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    page_number(sec.footer.paragraphs[0])

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(42)
r = p.add_run("EVENT WORKSPACE")
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor.from_string(ORANGE)
title = doc.add_paragraph(style="Title")
title.paragraph_format.space_before = Pt(5)
title.paragraph_format.space_after = Pt(10)
title.add_run("VSMS event operations,\nreporting & public page")
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(28)
run = subtitle.add_run("Implementation report  •  4 August 2026")
run.font.size = Pt(15)
run.font.color.rgb = RGBColor.from_string(MUTED)

summary = doc.add_table(rows=1, cols=1)
summary.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = summary.cell(0, 0)
shade(cell, INK)
cell.margin_top = Inches(0.22)
p = cell.paragraphs[0]
p.paragraph_format.space_after = Pt(8)
r = p.add_run("OUTCOME")
r.bold = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor.from_string(ORANGE)
p = cell.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("A complete, no-sidebar event workspace shaped around physical venue operations.")
r.bold = True
r.font.size = Pt(17)
r.font.color.rgb = RGBColor.from_string(WHITE)
p = cell.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("It combines Luma-inspired hierarchy with VSMS station availability, manpower planning, signups and attendance reporting, a safe public event landing page, and export-gated draft deletion.")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor.from_string("C7C5BD")

doc.add_paragraph()
table(doc, ["Delivery", "Result"], [
    ("Working branch", "nachiketh/eventsandmore — no new branch"),
    ("Source control", "Changes left unstaged and uncommitted, as requested"),
    ("Infrastructure", "Cloud-ready application changes only; no AWS or other deployment"),
    ("Validation", "Schema, OpenAPI, generated contract, lint, build, 33 backend tests, 28 static checks, queue regression, dependency audit, responsive browser QA"),
], [1.65, 4.8])

doc.add_page_break()
heading(doc, "Design translation", "Reference → VSMS")
body(doc, "The Luma reference supplied the hierarchy: concise navigation, event identity, tabbed workspace, overview actions, and a clear when-and-where summary. VSMS keeps that calm layering while replacing video-meeting language and guest-marketing patterns with physical venue operations.")
picture(doc, Path("/Users/nr/Downloads/image 2.png"), "Supplied Luma reference used for hierarchy and layering only.", 6.35)
bullet(doc, "Removed the global sidebar and its toggle; the product now uses a responsive 64px command bar.")
bullet(doc, "Physical venue, station order, availability slots, shifts, and assigned people are first-class facts.")
bullet(doc, "Warm operational canvas + near-black records + restrained orange status signal; no gradients, glass, glow, or generic bento styling.")

doc.add_page_break()
heading(doc, "Event register", "Operational overview")
body(doc, "The index is a chronological run sheet. Each event remains information-dense without becoming a dashboard mosaic: date rail, status, venue, signups, capacity, owner, and a direct 44px action.")
picture(doc, IMAGES / "events-list-desktop.png", "Desktop event register — chronological rail, compact physical-venue facts, truthful signup/capacity labels.", 6.55)
body(doc, "Pagination is consumed completely with a repeated-cursor guard, so search and upcoming/past filtering do not silently stop at the first page. The list intentionally says “Signups” rather than presenting active capacity as attendance.")

doc.add_page_break()
heading(doc, "Event workspace", "Plan and operate")
body(doc, "The authenticated workspace keeps event identity and primary actions stable while its tabs separate overview, attendees, operations, and settings.")
picture(doc, IMAGES / "event-workspace-overview-desktop.png", "Overview — event identity, real location and schedule, metrics, lifecycle, and audit context.", 6.55)

doc.add_page_break()
heading(doc, "Station and manpower plan", "Operations")
body(doc, "The operational view exposes the event’s real execution plan without inventing coverage: ordered stations, day/time availability, capacity, shifts, named assignments, and an aggregate manpower summary of filled positions, open positions, fully staffed shifts, and active placements.")
picture(doc, IMAGES / "event-workspace-operations-desktop.png", "Operations — ordered stations, day/time availability, capacity, shifts, and factual manpower assignments.", 5.9)

doc.add_page_break()
heading(doc, "Responsive operation", "Mobile evidence")
body(doc, "At phone widths the chronology rail, useful metadata, tab set, and primary actions remain intact. Labels condense before touch targets do; workspace tabs remain keyboard-operable and fit without page-level horizontal overflow.")
mobile_pair(doc, IMAGES / "events-list-mobile.png", "Event list", IMAGES / "event-workspace-mobile.png", "Event workspace")
table(doc, ["Accessibility baseline", "Implementation"], [
    ("Targets", "44px routine actions; 48px tabs and navigation controls"),
    ("Keyboard & semantics", "Visible focus; route-specific browser titles; ArrowLeft/ArrowRight/Home/End tabs; announced loading state; captioned attendee table with scoped columns"),
    ("Motion", "Transitions disabled when reduced motion is requested"),
    ("Layout", "No page-level horizontal overflow at tested phone and desktop viewports"),
], [1.75, 4.7])

doc.add_page_break()
heading(doc, "Public event landing page", "Safe by design")
body(doc, "The public page is a read-only venue briefing, available outside authentication only for non-draft events. It contains the event identity, physical location, dates, schedule, capacity, and status—never attendees, staff assignments, audit records, clinical data, or online-meeting details.")
picture(doc, IMAGES / "public-event-desktop.png", "Desktop public event page — physical venue briefing and schedule.", 6.0)
doc.add_page_break()
heading(doc, "Public and staff views", "Mobile evidence")
body(doc, "The public briefing and authenticated workspace retain distinct data boundaries at the same phone width.")
mobile_pair(doc, IMAGES / "public-event-mobile.png", "Mobile public page", IMAGES / "event-workspace-mobile.png", "Authenticated workspace")

doc.add_page_break()
heading(doc, "Implemented product surface", "Feature map")
table(doc, ["Area", "Delivered behavior"], [
    ("Shell & settings", "No sidebar; responsive command bar; theme, profile, command palette, and settings retained."),
    ("Overview", "Venue, schedule, event days, truthful signups/check-ins/attendance, lifecycle, and audit context."),
    ("Attendees", "Authenticated search and status filter with signed cursor pagination, 50-row loads, and explicit loading/error/empty states."),
    ("Operations", "Ordered stations, per-day availability, shifts, assigned people/roles, and truthful aggregate manpower coverage."),
    ("Settings", "Edit link, lifecycle controls, cancellation, export, and guarded draft deletion in one workspace."),
    ("Public landing", "Read-only public projection for non-draft events, surfaced as a first-class event action."),
    ("API contract", "Canonical metrics, attendees, export, delete, public event, GET, and PATCH paths documented and generated."),
], [1.45, 5.0])

doc.add_heading("Files and boundaries", level=2)
body(doc, "Frontend changes live in the existing React application shell and event feature. Backend changes reuse the existing Express, Zod, Prisma, OpenAPI, rate-limit, authentication, and authorization patterns. No new dependency or speculative platform abstraction was added.")
body(doc, "Teammate-owned station import/configuration, participant, QR, screening, and registration implementations were not replaced. The station/manpower planning UX remains visible for review, while three persistence operations fail explicitly with documented 501 responses until the teammate-owned station model work lands.")

doc.add_page_break()
heading(doc, "Security and deletion model", "High-impact controls")
table(doc, ["Boundary", "Control"], [
    ("Public read", "UUID validation; 120 requests/minute; narrow projection; drafts use the same 404 as unknown IDs; cache max 60 seconds."),
    ("Private reporting", "Authenticated; manager/admin authority; no-store; 60 requests/minute per user; attendee pages capped at 100."),
    ("Export", "Versioned JSON with event, aggregate metrics, days, availability, stations, shifts, essential staffing fields, and narrow attendee rows; free-form assignment notes are excluded."),
    ("Receipt", "Domain-separated HMAC-SHA256; base64url; 15-minute expiry; binds event, version, actor, and stable export hash."),
    ("Delete", "Serializable transaction; exact event name/version/receipt; only an unpopulated DRAFT with zero registrations and consents."),
    ("Retention", "Global users, participants, station templates, and event audit rows are preserved; a general audit tombstone records deletion."),
], [1.35, 5.1])

doc.add_heading("Session and provider hardening", level=2)
table(doc, ["Risk", "Implemented control"], [
    ("Browser credentials", "Cognito access and refresh tokens remain Secure + HttpOnly; no credential token is exposed to JavaScript."),
    ("CSRF", "Rotating double-submit token, allowed-Origin check, Fetch Metadata check, and protection on cookie-authenticated mutations."),
    ("Compatibility", "Existing signed Bearer-token API clients remain supported without weakening the browser cookie path."),
    ("Key rotation", "Unknown Cognito signing-key IDs trigger one immediate JWKS refresh; JWKS fetches time out after five seconds."),
    ("Provider stalls", "Cognito token and account requests time out after ten seconds instead of holding API requests indefinitely."),
    ("Stale browser state", "Without the authoritative session/CSRF cookie, cached profile state is cleared instead of presenting a false authenticated session."),
    ("Blocked Web Storage", "Session state fails closed; device identification uses an in-memory UUID fallback rather than breaking requests."),
    ("Return path", "Post-login redirects reject external, protocol-relative, backslash, encoded-control, and malformed paths."),
], [1.35, 5.1])

doc.add_heading("Why deletion is deliberately narrow", level=2)
body(doc, "Populated or clinical events cannot be deleted; export alone does not make a purge safe. Only an empty DRAFT can be removed after a fresh receipt. The public page is read-only and rate-limited, so CAPTCHA is deferred until a public mutation exists.")

heading(doc, "Continuation hardening", "4 August 2026")
table(doc, ["Finding", "Resolution and direct impact"], [
    ("API startup crash", "Removed unused imports introduced in Keefe’s participant-QR merge, including undeclared js-yaml/Socket.IO references, and removed the undefined io export. The API now loads and listens normally."),
    ("Session mismatch", "Completed the repository’s documented HttpOnly Cognito-cookie + CSRF design. Previously the browser could not authenticate event calls and refresh sent invalid strict-JSON null."),
    ("Reload CSRF race", "Mutation requests now recover the rotating CSRF token from its readable cookie when in-memory state has not hydrated after a page reload."),
    ("Local port collision", "Aligned backend, example configuration, Vite proxy, and OpenAPI on 5050; avoids macOS Control Center’s common port-5000 listener."),
    ("Setup drift", "README, HTTPS smoke test, Bruno environment, and API collection now agree on managed Cognito login and the 5173/5050 local HTTPS surface."),
    ("Audit pagination", "Uses timestamp plus record ID, so equal-timestamp audit rows are not skipped between signed-cursor pages."),
    ("Lifecycle regression", "Corrected Keefe’s merge-adjusted event commands from the legacy id field to the schema’s eventId key, retained canManage in responses, and covered publish/cancel with executable tests."),
    ("Event plan integrity", "Event creation persists selected event days; create/update reject day slots outside the overall schedule; simple physical-location edits no longer depend on an absent station model; unsupported station/staff-plan writes fail explicitly with 501."),
    ("Versioned staffing", "Staff removal now validates the real assignment ID and atomically checks/increments the event version before deletion, preventing stale editors from overwriting each other."),
    ("Shared route pass", "Event and screening endpoints now share one authentication and mutation-limit pass, eliminating duplicate user lookups and double-counted screening writes."),
    ("Scoped operations", "Screening access is event-scoped for managers and station-scoped for screeners; clinical aggregates and full event-management actions require event-level management authority."),
    ("Privacy projection", "Exports omit free-form assignment notes; ordinary assigned staff receive only roster user ID/display name and never contact, role-status metadata, or private notes."),
    ("Audit boundary", "Event mutations retain real request ID, IP, and device context; malformed request/device IDs cannot reach PostgreSQL UUID columns."),
    ("Transport headers", "Live API policy denies framing, camera, microphone, and geolocation; Swagger-only inline CSP allowances are development-only and smoke-tested."),
    ("Log and seed safety", "Removed redundant raw validation logging; production blocks demonstration seeding and seed logs no longer publish the fixed demo QR pass token."),
    ("Storage reliability", "Blocked Web Storage no longer crashes app startup, theme changes, session restoration, or device identification; sensitive state still fails closed."),
    ("Supply-chain claim", "Removed two mismatched/unused runtime signature prototypes that could silently skip verification. Signed immutable artifacts remain an explicit deployment-pipeline requirement, not an unproved application claim."),
    ("Error disclosure", "Unknown failures never expose raw messages. Existing bounded status hints retain correct 4xx/5xx semantics with generic copy; unclassified failures become generic 500 responses."),
    ("Role-consistent actions", "The event register now exposes creation only to verified administrator/event-manager grants, matching the shell and backend instead of treating every non-staff role as privileged."),
    ("Theme interaction", "Removed the nonessential native view-transition layer after it dropped rapid repeat clicks; each theme choice now applies immediately with the existing icon feedback."),
    ("Landing evidence", "Replaced the preview placeholder with a real, PII-free workspace screenshot, truthful captions, and a connected-workflow claim instead of unsupported offline-sync copy."),
    ("Dependency exposure", "Updated ip-address to 10.4.0 and removed unused Geist/Next/sharp graph. Backend production audit is clean."),
    ("Bundle size", "Route-level lazy loading, including the authenticated shell and event register, reduced initial JavaScript from 818.37 kB to 294.47 kB — a 64.0% reduction and no size warning."),
], [1.55, 4.9])
body(doc, "The frontend audit still reports the React Router RSC advisory GHSA-qwww-vcr4-c8h2. This Vite BrowserRouter application does not import or use the affected unstable RSC APIs, and the advisory’s patched 8.3.0 release is not published in the installed package line. The non-reachable advisory is recorded rather than forcing an incompatible or nonexistent upgrade.")

heading(doc, "Verification record", "4 August 2026")
table(doc, ["Gate", "Outcome"], [
    ("Prisma schema", "Valid"),
    ("OpenAPI", "Valid; four warnings: one existing auth redirect and three documented 501-only teammate stubs"),
    ("Generated client", "Matches canonical OpenAPI contract"),
    ("Frontend lint", "Passed"),
    ("Production build", "Passed; initial JavaScript 294.47 kB (96.38 kB gzip), no size warning"),
    ("Backend unit/security tests", "33/33 passed; includes event/station-scoped screening and clinical-metrics authorization"),
    ("Static integration contracts", "28/28 passed; includes role-consistent actions, reliable theme toggling, storage fallback, production seed guard, release-integrity claim, shared route-pass, CSRF, schema, and validation checks"),
    ("Database integration", "Guarded harness stopped before migrations: local vsms_dev_test is absent and the configured role cannot create databases; no cloud database was changed"),
    ("Production dependency audit", "Backend: 0 vulnerabilities. Frontend: two RSC-only React Router advisory instances, non-reachable here"),
    ("Queue regression", "1/1 passed"),
    ("Impeccable detector", "No findings"),
    ("Browser QA", "Desktop/mobile evidence retained; local landing and real product preview rendered without placeholder content; authenticated live QA is environment-gated"),
    ("HTTPS smoke", "Passed: trusted frontend/API, proxied private-route 401, plain HTTP rejection, anti-framing CSP, and device-permission denial"),
    ("Whitespace", "git diff --check passed"),
], [2.0, 4.45])

doc.add_heading("Known pre-existing / teammate-owned blockers", level=2)
body(doc, "The wider event service still contains three documented station import/configuration/assignment 501 paths pending teammate PR #65. They were not replaced because that implementation remains teammate-owned and changing it here would collide with the active station model work.")
body(doc, "The Sol Advisor native-agent preflight also reports two missing local template files: sol-advisor-terra-implementer.toml and sol-advisor-sol-reviewer.toml. The continuation therefore used bounded local verification rather than silently substituting another delegation path.")

heading(doc, "Handoff", "Ready for review")
body(doc, "The requested MVP is implemented on the existing branch and left uncommitted. It can be reviewed locally using the captured desktop and mobile evidence in .codex/images. No cloud resource was created, restarted, or deployed.")
table(doc, ["Artifact", "Location"], [
    ("Design system", str(ROOT / "DESIGN.md")),
    ("Implementation report", str(OUTPUT)),
    ("Screenshot evidence", str(IMAGES)),
    ("Security note", "/Users/nr/documents/Obsidian Vault/Secure Coding/wiki/vsms-event-workspace-security.md"),
], [1.55, 4.9])
body(doc, "Recommended next step: have a local database administrator create vsms_dev_test, rerun the guarded integration harness, then integrate the teammate-owned station 501 implementations before any deployment.")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
